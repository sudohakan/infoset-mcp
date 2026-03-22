#!/usr/bin/env python3
"""
Work Sync DOCX Report Generator
Usage: echo '{json}' | python3 generate-report.py /path/to/output.docx
"""

import sys
import json
import os

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import docx.opc.constants
except ImportError:
    print("ERROR: python-docx is not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(2)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_bg(cell, hex_color: str):
    """Set table cell background color via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_page_number(section):
    """Add page number to footer (center-aligned)."""
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    para.add_run(' / ')
    run2 = para.add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    instrText2 = OxmlElement('w:instrText')
    instrText2.text = 'NUMPAGES'
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    run2._r.append(fldChar3)
    run2._r.append(instrText2)
    run2._r.append(fldChar4)


def add_page_break(doc):
    """Add a page break paragraph."""
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)
    return para


def style_heading(para, level: int):
    """Apply heading style to paragraph."""
    style_name = f'Heading {level}'
    try:
        para.style = style_name
    except Exception:
        pass


def add_section_heading(doc, text: str, level: int = 1):
    """Add a section heading, with page break before Heading 1."""
    if level == 1:
        add_page_break(doc)
    para = doc.add_heading(text, level=level)
    return para


def make_table_header(table, headers: list, bg_hex: str = '2E4057'):
    """Style the first row of a table as a header."""
    hdr_row = table.rows[0]
    for i, cell in enumerate(hdr_row.cells):
        cell.text = headers[i] if i < len(headers) else ''
        set_cell_bg(cell, bg_hex)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0] if para.runs else para.add_run(cell.text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'


def set_cell_text(cell, text: str, bold: bool = False, font_size: int = 10,
                  align=WD_ALIGN_PARAGRAPH.LEFT):
    """Set cell text with formatting."""
    cell.text = str(text)
    para = cell.paragraphs[0]
    para.alignment = align
    if para.runs:
        run = para.runs[0]
        run.bold = bold
        run.font.size = Pt(font_size)
        run.font.name = 'Calibri'


def tier_row_bg(tier: int) -> str | None:
    if tier == 1:
        return 'FFC8C8'  # light red
    elif tier == 2:
        return 'FFE6C8'  # light orange
    return None


# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------

def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    # A4 landscape
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    add_page_number(section)

    # Default paragraph style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    return doc


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def build_cover(doc: Document, data: dict):
    """Cover / Title section."""
    summary = data.get('summary', {})
    date_str = data.get('date', '')

    doc.add_paragraph()
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run('Work Sync Raporu')
    run.bold = True
    run.font.size = Pt(28)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = date_para.add_run(date_str)
    run2.font.size = Pt(16)
    run2.font.name = 'Calibri'
    run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    summary_para = doc.add_paragraph()
    summary_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary_text = (
        f"Infoset: {summary.get('infosetTickets', 0)} ticket  |  "
        f"DevOps: {summary.get('devopsTasks', 0)} görev  |  "
        f"Eşleşme: {summary.get('matched', 0)}  |  "
        f"İş Planı: {summary.get('workPlanItems', 0)}  |  "
        f"Kapasite Bu Hafta: {summary.get('capacityThisWeek', '-')}  |  "
        f"Gelecek Hafta: {summary.get('capacityNextWeek', '-')}"
    )
    run3 = summary_para.add_run(summary_text)
    run3.font.size = Pt(12)
    run3.font.name = 'Calibri'
    run3.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def build_genel_bakis(doc: Document, data: dict):
    """Bölüm 1: Genel Bakış"""
    add_section_heading(doc, 'Bölüm 1: Genel Bakış', level=1)

    summary = data.get('summary', {})

    add_section_heading(doc, 'Kaynak Dağılımı', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    make_table_header(table, ['Kaynak', 'Adet'])
    rows_data = [
        ('Infoset Ticket', summary.get('infosetTickets', 0)),
        ('DevOps Görev', summary.get('devopsTasks', 0)),
        ('Eşleşen', summary.get('matched', 0)),
        ('İş Planı Kalemi', summary.get('workPlanItems', 0)),
    ]
    for label, val in rows_data:
        row = table.add_row()
        set_cell_text(row.cells[0], label)
        set_cell_text(row.cells[1], str(val), align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    add_section_heading(doc, 'Kapasite Özeti', level=2)
    cap_table = doc.add_table(rows=1, cols=2)
    cap_table.style = 'Table Grid'
    make_table_header(cap_table, ['Dönem', 'Kullanım'])
    cap_rows = [
        ('Bu Hafta', summary.get('capacityThisWeek', '-')),
        ('Gelecek Hafta', summary.get('capacityNextWeek', '-')),
    ]
    for label, val in cap_rows:
        row = cap_table.add_row()
        set_cell_text(row.cells[0], label)
        set_cell_text(row.cells[1], str(val), align=WD_ALIGN_PARAGRAPH.CENTER)


def build_eslestirme(doc: Document, data: dict):
    """Bölüm 2: Eşleştirme Raporu"""
    add_section_heading(doc, 'Bölüm 2: Eşleştirme Raporu', level=1)

    matching = data.get('matching', {})

    # Matched pairs
    add_section_heading(doc, 'Eşleşen Çiftler', level=2)
    matched = matching.get('matched', [])
    if matched:
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        make_table_header(table, ['Infoset ID', 'DevOps ID(s)', 'Müşteri', 'Infoset Konu', 'DevOps Başlık'])
        for item in matched:
            row = table.add_row()
            devops_ids = ', '.join(str(i) for i in item.get('devopsIds', []))
            set_cell_text(row.cells[0], str(item.get('infosetId', '')), align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(row.cells[1], devops_ids, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(row.cells[2], item.get('customer', ''))
            set_cell_text(row.cells[3], item.get('infosetSubject', ''))
            set_cell_text(row.cells[4], item.get('devopsTitle', ''))
    else:
        doc.add_paragraph('Eşleşen çift bulunamadı.')

    doc.add_paragraph()

    # Infoset only
    add_section_heading(doc, 'Sadece Infoset (DevOps Karşılığı Yok)', level=2)
    infoset_only = matching.get('infosetOnly', [])
    if infoset_only:
        table2 = doc.add_table(rows=1, cols=3)
        table2.style = 'Table Grid'
        make_table_header(table2, ['ID', 'Müşteri', 'Konu'])
        for item in infoset_only:
            row = table2.add_row()
            set_cell_text(row.cells[0], str(item.get('id', '')), align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(row.cells[1], item.get('customer', ''))
            set_cell_text(row.cells[2], item.get('subject', ''))
    else:
        doc.add_paragraph('Yok.')

    doc.add_paragraph()

    # DevOps only
    add_section_heading(doc, 'Sadece DevOps (Infoset Karşılığı Yok)', level=2)
    devops_only = matching.get('devopsOnly', [])
    if devops_only:
        table3 = doc.add_table(rows=1, cols=4)
        table3.style = 'Table Grid'
        make_table_header(table3, ['ID', 'Başlık', 'Sprint', 'Durum'])
        for item in devops_only:
            row = table3.add_row()
            set_cell_text(row.cells[0], str(item.get('id', '')), align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(row.cells[1], item.get('title', ''))
            set_cell_text(row.cells[2], item.get('sprint', ''))
            set_cell_text(row.cells[3], item.get('state', ''))
    else:
        doc.add_paragraph('Yok.')


def build_oncelik(doc: Document, data: dict):
    """Bölüm 3: Öncelik Analizi"""
    add_section_heading(doc, 'Bölüm 3: Öncelik Analizi', level=1)

    tiers = data.get('tiers', {})
    headers = ['#', 'Başlık', 'Kaynak', 'Durum', 'Skor', 'Efor', 'Müşteri', 'Sprint', 'Bekleyen']

    tier_labels = {1: 'Tier 1 — Acil (skor 70+)', 2: 'Tier 2 — Bu Hafta (skor 50-69)', 3: 'Tier 3 — Backlog (skor <50)'}

    for tier_num in [1, 2, 3]:
        tier_key = f'tier{tier_num}'
        items = tiers.get(tier_key, [])
        add_section_heading(doc, tier_labels.get(tier_num, f'Tier {tier_num}'), level=2)

        if not items:
            doc.add_paragraph(f'Tier {tier_num} kalemi yok.')
            continue

        for item in items:
            # Item header with score and title
            emoji = item.get('emoji', '')
            score = item.get('score', 0)
            title = item.get('title', '')
            item_id = item.get('id', '')
            header_text = f"{emoji} [{score}] {title}"

            p_header = doc.add_paragraph()
            run = p_header.add_run(header_text)
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            bg = tier_row_bg(tier_num)
            if bg:
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn as _qn
                shd = OxmlElement('w:shd')
                shd.set(_qn('w:fill'), bg)
                shd.set(_qn('w:val'), 'clear')
                p_header.paragraph_format.element.get_or_add_pPr().append(shd)

            # Metadata line
            meta_parts = []
            if item.get('source'):
                meta_parts.append(f"Kaynak: {item['source']}")
            if item.get('customer'):
                meta_parts.append(f"Müşteri: {item['customer']}")
            if item.get('state'):
                meta_parts.append(f"Durum: {item['state']}")
            if item.get('sprint') and item['sprint'] != '—':
                meta_parts.append(f"Sprint: {item['sprint']}")
            if item.get('effort'):
                meta_parts.append(f"Efor: {item['effort']}s")
            if item.get('waiting'):
                meta_parts.append(f"Bekleyen: {item['waiting']}")
            if meta_parts:
                p_meta = doc.add_paragraph()
                run_meta = p_meta.add_run(' | '.join(meta_parts))
                run_meta.font.name = 'Calibri'
                run_meta.font.size = Pt(9)
                run_meta.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            # Action summary — the AI analysis of what the problem is and what to do
            action = item.get('actionSummary', '')
            if action:
                p_action = doc.add_paragraph()
                run_action = p_action.add_run(action)
                run_action.font.name = 'Calibri'
                run_action.font.size = Pt(10)
                run_action.italic = True

            # Notes — additional findings
            notes = item.get('notes', '')
            if notes:
                p_notes = doc.add_paragraph()
                run_notes = p_notes.add_run(notes)
                run_notes.font.name = 'Calibri'
                run_notes.font.size = Pt(9)
                run_notes.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

            # Separator
            doc.add_paragraph('─' * 80).runs[0].font.size = Pt(6)

        doc.add_paragraph()


def build_musteri(doc: Document, data: dict):
    """Bölüm 4: Müşteri Bazlı Çapraz Görünüm"""
    add_section_heading(doc, 'Bölüm 4: Müşteri Bazlı Çapraz Görünüm', level=1)

    customer_view = data.get('customerView', {})
    if not customer_view:
        doc.add_paragraph('Müşteri verisi bulunamadı.')
        return

    for customer, cdata in customer_view.items():
        add_section_heading(doc, customer, level=2)

        # Summary line
        summary_para = doc.add_paragraph()
        run = summary_para.add_run(
            f"Toplam Efor: {cdata.get('totalEffort', 0)} saat  |  "
            f"En Yüksek Skor: {cdata.get('highestScore', 0)}"
        )
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)

        # Infoset tickets
        infoset_items = cdata.get('infoset', [])
        if infoset_items:
            doc.add_paragraph('Infoset Ticket\'ları:').runs[0].font.name = 'Calibri'
            it_table = doc.add_table(rows=1, cols=4)
            it_table.style = 'Table Grid'
            make_table_header(it_table, ['ID', 'Konu', 'Kategori', 'Yaş (gün)'])
            for it in infoset_items:
                row = it_table.add_row()
                set_cell_text(row.cells[0], str(it.get('id', '')), align=WD_ALIGN_PARAGRAPH.CENTER)
                set_cell_text(row.cells[1], it.get('subject', ''))
                set_cell_text(row.cells[2], it.get('category', ''))
                set_cell_text(row.cells[3], str(it.get('ageDays', '')), align=WD_ALIGN_PARAGRAPH.CENTER)
                # Action summary row spanning all columns
                action = it.get('actionSummary', '')
                if action:
                    detail_row = it_table.add_row()
                    merged = detail_row.cells[0].merge(detail_row.cells[3])
                    run = merged.paragraphs[0].add_run(f"→ {action}")
                    run.font.name = 'Calibri'
                    run.font.size = Pt(9)
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            doc.add_paragraph()

        # DevOps tasks
        devops_items = cdata.get('devops', [])
        if devops_items:
            doc.add_paragraph('DevOps Görevleri:').runs[0].font.name = 'Calibri'
            do_table = doc.add_table(rows=1, cols=4)
            do_table.style = 'Table Grid'
            make_table_header(do_table, ['ID', 'Başlık', 'Durum', 'Sprint'])
            for di in devops_items:
                row = do_table.add_row()
                set_cell_text(row.cells[0], str(di.get('id', '')), align=WD_ALIGN_PARAGRAPH.CENTER)
                set_cell_text(row.cells[1], di.get('title', ''))
                set_cell_text(row.cells[2], di.get('state', ''))
                set_cell_text(row.cells[3], di.get('sprint', ''), align=WD_ALIGN_PARAGRAPH.CENTER)
                # Action summary row spanning all columns
                action = di.get('actionSummary', '')
                if action:
                    detail_row = do_table.add_row()
                    merged = detail_row.cells[0].merge(detail_row.cells[3])
                    run = merged.paragraphs[0].add_run(f"→ {action}")
                    run.font.name = 'Calibri'
                    run.font.size = Pt(9)
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            doc.add_paragraph()


def build_kapasite(doc: Document, data: dict):
    """Bölüm 5: Haftalık Kapasite Planı"""
    add_section_heading(doc, 'Bölüm 5: Haftalık Kapasite Planı', level=1)

    capacity = data.get('capacity', {})
    weeks = capacity.get('weeks', [])

    if weeks:
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        make_table_header(table, ['Hafta', 'Mevcut (gün)', 'Planlanan (gün)', 'Kullanım'])
        for week in weeks:
            row = table.add_row()
            util_str = str(week.get('utilization', ''))
            set_cell_text(row.cells[0], week.get('label', ''))
            set_cell_text(row.cells[1], str(week.get('available', '')), align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(row.cells[2], str(week.get('planned', '')), align=WD_ALIGN_PARAGRAPH.CENTER)
            # Color code high utilization
            set_cell_text(row.cells[3], util_str, align=WD_ALIGN_PARAGRAPH.CENTER)
            try:
                util_num = float(util_str.replace('%', ''))
                if util_num > 100:
                    set_cell_bg(row.cells[3], 'FFC8C8')
                elif util_num > 90:
                    set_cell_bg(row.cells[3], 'FFE6C8')
            except (ValueError, AttributeError):
                pass

    doc.add_paragraph()
    totals_para = doc.add_paragraph()
    run = totals_para.add_run(
        f"Toplam İş: {capacity.get('totalWork', 0)} gün  |  "
        f"Haftalık Kapasite: {capacity.get('weeklyCapacity', 0)} gün"
    )
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)


def build_aging(doc: Document, data: dict):
    """Bölüm 6: Trend / Aging Raporu"""
    add_section_heading(doc, 'Bölüm 6: Trend / Aging Raporu', level=1)

    aging = data.get('aging', {})

    add_section_heading(doc, 'Sprint Carry-Over', level=2)
    sprint_carry = aging.get('sprintCarry', [])
    if sprint_carry:
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        make_table_header(table, ['Sprint', 'Adet', 'Not'])
        for item in sprint_carry:
            row = table.add_row()
            set_cell_text(row.cells[0], item.get('sprint', ''))
            set_cell_text(row.cells[1], str(item.get('count', '')), align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(row.cells[2], item.get('note', ''))
    else:
        doc.add_paragraph('Sprint carry-over yok.')

    doc.add_paragraph()

    add_section_heading(doc, 'Yaş Dağılımı', level=2)
    dist = aging.get('distribution', {})
    if dist:
        table2 = doc.add_table(rows=1, cols=2)
        table2.style = 'Table Grid'
        make_table_header(table2, ['Yaş Aralığı', 'Adet'])
        dist_rows = [
            ('60+ gün', dist.get('over60', 0)),
            ('30-60 gün', dist.get('30to60', 0)),
            ('14-30 gün', dist.get('14to30', 0)),
            ('14 günden az', dist.get('under14', 0)),
        ]
        for label, val in dist_rows:
            row = table2.add_row()
            set_cell_text(row.cells[0], label)
            set_cell_text(row.cells[1], str(val), align=WD_ALIGN_PARAGRAPH.CENTER)
    else:
        doc.add_paragraph('Yaş dağılımı verisi yok.')


def build_duplicates(doc: Document, data: dict):
    """Bölüm 7: Duplicate / İlişkili İşler"""
    add_section_heading(doc, 'Bölüm 7: Duplicate / İlişkili İşler', level=1)

    duplicates = data.get('duplicates', [])
    if not duplicates:
        doc.add_paragraph('Duplicate bulunamadı.')
        return

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    make_table_header(table, ['ID\'ler', 'Neden', 'Aksiyon'])
    for item in duplicates:
        row = table.add_row()
        ids_str = ', '.join(str(i) for i in item.get('ids', []))
        set_cell_text(row.cells[0], ids_str, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[1], item.get('reason', ''))
        set_cell_text(row.cells[2], item.get('action', ''))


def build_aksiyonlar(doc: Document, data: dict):
    """Bölüm 8: Aksiyon Önerileri"""
    add_section_heading(doc, 'Bölüm 8: Aksiyon Önerileri', level=1)

    actions = data.get('actions', [])
    if not actions:
        doc.add_paragraph('Aksiyon önerisi yok.')
        return

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    make_table_header(table, ['Öncelik', 'Aksiyon'])

    priority_colors = {
        'BUGÜN': 'FFC8C8',
        'BU HAFTA': 'FFE6C8',
        'EKSİK': 'FFF3C8',
    }

    for item in actions:
        row = table.add_row()
        priority = item.get('priority', '')
        set_cell_text(row.cells[0], priority, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[1], item.get('text', ''))
        bg = priority_colors.get(priority.upper(), None)
        if bg:
            set_cell_bg(row.cells[0], bg)
            set_cell_bg(row.cells[1], bg)


def build_sync_sonuclari(doc: Document, data: dict):
    """Bölüm 9: Sync Sonuçları"""
    add_section_heading(doc, 'Bölüm 9: Sync Sonuçları', level=1)

    sync = data.get('syncResults', {})
    if not sync:
        doc.add_paragraph('Sync sonucu verisi yok.')
        return

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    make_table_header(table, ['İşlem', 'Infoset', 'DevOps', 'İş Planı', 'Takvim'])

    operations = [
        ('Yeni', 'new'),
        ('Güncellendi', 'updated'),
        ('Kapatıldı', 'closed'),
    ]
    for label, key in operations:
        row = table.add_row()
        op_data = sync.get(key, {})
        set_cell_text(row.cells[0], label, bold=True)
        set_cell_text(row.cells[1], str(op_data.get('infoset', 0)), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[2], str(op_data.get('devops', 0)), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[3], str(op_data.get('workPlan', 0)), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[4], str(op_data.get('calendar', 0)), align=WD_ALIGN_PARAGRAPH.CENTER)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    if len(sys.argv) < 2:
        print("Usage: echo '{json}' | python3 generate-report.py /path/to/output.docx", file=sys.stderr)
        sys.exit(1)

    output_path = sys.argv[1]

    # Read JSON from stdin
    try:
        raw = sys.stdin.read()
        data = json.loads(raw)
    except json.JSONDecodeError as e:
        print(f"ERROR: Failed to parse JSON input: {e}", file=sys.stderr)
        sys.exit(1)

    # Ensure output directory exists
    output_dir = os.path.dirname(os.path.abspath(output_path))
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    # Build document
    doc = setup_document()

    build_cover(doc, data)
    build_genel_bakis(doc, data)
    build_eslestirme(doc, data)
    build_oncelik(doc, data)
    build_musteri(doc, data)
    build_kapasite(doc, data)
    build_aging(doc, data)
    build_duplicates(doc, data)
    build_aksiyonlar(doc, data)
    build_sync_sonuclari(doc, data)

    doc.save(output_path)
    print(f"Report saved: {output_path}")


if __name__ == '__main__':
    main()
